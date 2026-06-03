import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Title
title = doc.add_heading('Purplle Store Intelligence - Execution Guide', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Follow these instructions to set up and run the Purplle Store Intelligence System.')

# Note about Videos
doc.add_heading('1. Important Note Regarding Videos', level=1)
p1 = doc.add_paragraph()
p1.add_run('Please note: ').bold = True
p1.add_run('Due to GitHub file size limits and data privacy constraints, the CCTV video feeds are ')
p1.add_run('NOT').bold = True
p1.add_run(' included in this repository submission.')
doc.add_paragraph('To run the pipeline with actual video feeds, please ensure you place your CCTV video files (.mp4 or .avi) inside the following directory:')
doc.add_paragraph('data/videos/', style='Intense Quote')
doc.add_paragraph('The system requires these videos to be present for the computer vision ingestion pipeline to function properly.')

# Setup
doc.add_heading('2. Environment Setup', level=1)
doc.add_paragraph('You can set up the environment using the provided Makefile or manually.')
doc.add_heading('Using Makefile (Recommended):', level=2)
doc.add_paragraph('make setup', style='Intense Quote')
doc.add_paragraph('Activate the virtual environment:')
doc.add_paragraph(r'.\.venv\Scripts\activate  # On Windows', style='Intense Quote')
doc.add_paragraph('source .venv/bin/activate  # On Linux/Mac', style='Intense Quote')

# Running the system
doc.add_heading('3. Running the System', level=1)
doc.add_paragraph('The system consists of a FastAPI backend and a Streamlit dashboard. You will need to open two separate terminal windows.')

doc.add_heading('Terminal 1: Start the Backend (FastAPI)', level=2)
doc.add_paragraph('Run the API server:')
doc.add_paragraph('make run', style='Intense Quote')
doc.add_paragraph('The backend will be available at: http://localhost:8000')

doc.add_heading('Terminal 2: Start the Dashboard (Streamlit)', level=2)
doc.add_paragraph('Run the analytics dashboard:')
doc.add_paragraph('make run-dashboard', style='Intense Quote')
doc.add_paragraph('The dashboard will open automatically in your browser at: http://localhost:8501')

doc.add_heading('Running Data Simulation / Computer Vision pipeline', level=2)
doc.add_paragraph('If you want to mock the data stream without video ingestion:')
doc.add_paragraph('make stream-mock', style='Intense Quote')
doc.add_paragraph('If you want to run the actual computer vision pipeline on the videos in data/videos/:')
doc.add_paragraph('python pipeline/main.py', style='Intense Quote')

# Save
doc.save('RUNNING_INSTRUCTIONS.docx')
print("Document saved successfully.")
